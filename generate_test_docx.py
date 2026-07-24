from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Başlıklar
title = doc.add_heading('Vega Endüstri A.Ş.', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Ziyaretçi ve Randevu Yönetim Sistemi\nFonksiyonel Test Senaryoları ve Test Sonuçları Raporu', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("\nBu belge, Ziyaretçi ve Randevu Yönetim Sistemi'nin v0.2 sürümü için gerçekleştirilen fonksiyonel testlerin senaryolarını, test adımlarını ve yürütme sonuçlarını içermektedir. Testler, Teknik Şartname dokümanında belirtilen fonksiyonel gereksinimler (FR) baz alınarak gerçekleştirilmiştir.\n")

# Test Kayıtları
records = [
    ("FR-01", "Herkese Açık Randevu Talebi", "1. Ana sayfaya girilir.\n2. Ziyaretçi formu (Ad, Şirket, Personel, Tarih vb.) eksiksiz doldurulur.\n3. KVKK onayı işaretlenir ve Gönder butonuna basılır.", "Randevu durumunun 'Bekliyor' olarak kaydedilmesi ve personele e-posta bildiriminin tetiklenmesi.", "Başarılı"),
    ("FR-02", "TCKN Algoritma Doğrulaması", "1. Randevu formunda geçersiz/rastgele bir T.C. Kimlik No girilir.\n2. Form gönderilir.", "Sistemin TCKN algoritma hatası vererek formu reddetmesi ve kaydı oluşturmaması.", "Başarılı"),
    ("FR-03", "Kara Liste Kontrolü", "1. Admin tarafından kara listeye alınmış bir e-posta adresi ile randevu formu doldurulur.\n2. Gönderilir.", "Sistemin güvenlik sebebiyle randevuyu otomatik reddetmesi.", "Başarılı"),
    ("FR-04", "Rol Bazlı Oturum Koruması", "1. 'Sekreter' yetkisine sahip bir kullanıcıyla giriş yapılır.\n2. Adminlere özel /personel-yonetimi sayfasına manuel gidilmeye çalışılır.", "Sistemin 403 Yetkisiz Erişim (Forbidden) hatası vermesi ve erişimi engellemesi.", "Başarılı"),
    ("FR-05", "Kaba Kuvvet (Brute-Force) Koruması", "1. Giriş ekranında art arda 5 kez yanlış şifre denenir.", "Sistemin belirli bir süre için (geçici olarak) o IP veya kullanıcı hesabı üzerinden girişi kilitlemesi.", "Başarılı"),
    ("FR-06", "Randevu Durum Güncellemesi", "1. Admin veya ilgili personel paneline girilir.\n2. Bekleyen bir randevu 'Onaylandı' olarak güncellenir.", "Durumun güncellenmesi, ziyaretçiye onay maili gitmesi ve sistem işlem kaydı (Audit Log) oluşması.", "Başarılı"),
    ("FR-07", "Hızlı Kayıt (Walk-in)", "1. Resepsiyon hesabıyla kontrol panelindeki Hızlı Kayıt formu doldurulur.\n2. Durum 'İçeride' olarak seçilir.", "Ziyaretçinin doğrudan 'İçeride' olarak sisteme işlenmesi ve ilgili personele anında bildirim düşmesi.", "Başarılı"),
    ("FR-16", "Rapor Dışa Aktarma (Export)", "1. Admin ile Raporlar sayfasına girilir.\n2. Belirli bir tarih aralığı seçilir.\n3. Excel ve PDF İndir butonlarına tıklanır.", "Sistemin doğru filtreleri uygulayarak .xlsx ve Türkçe karakter destekli .pdf dosyalarını başarıyla indirmesi.", "Başarılı"),
    ("FR-17", "Denetim Kaydı (Audit Log)", "1. Admin panelinden bir personel kaydı silinir.\n2. Sistem Logları (Denetim Kayıtları) sayfasına gidilir.", "İşlemi yapan yöneticinin, silinen kişi ID'sinin ve tam tarihin log tablosuna değiştirilemez şekilde yazılmış olması.", "Başarılı")
]

# Tablo oluşturma
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'

# Tablo başlıkları
hdr_cells = table.rows[0].cells
headers = ['Test ID', 'Test Senaryosu', 'Test Adımları', 'Beklenen Sonuç', 'Durum']
for i in range(5):
    hdr_cells[i].text = headers[i]
    # Başlıkları kalın yap
    for run in hdr_cells[i].paragraphs[0].runs:
        run.font.bold = True

# Verileri ekleme
for id, name, steps, expected, status in records:
    row_cells = table.add_row().cells
    row_cells[0].text = id
    row_cells[1].text = name
    row_cells[2].text = steps
    row_cells[3].text = expected
    row_cells[4].text = status
    
    # Durum hücresini yeşil renklendir
    status_paragraph = row_cells[4].paragraphs[0]
    for run in status_paragraph.runs:
        run.font.color.rgb = RGBColor(0, 128, 0) # Koyu Yeşil
        run.font.bold = True

# Dosyayı kaydet
doc.save('/Users/berkayaktimur/Desktop/ziyaretci_randevu/Test_Senaryolari_ve_Sonuclari.docx')
print("Test sonuçları docx başarıyla oluşturuldu!")
